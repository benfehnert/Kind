#!/usr/bin/env python3
"""Convert mobile-deterministic-algorithms.md to a Word document."""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def add_rich_text(paragraph, text):
    """Parse **bold** and `code` inline markers."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        chunk = match.group(0)
        if chunk.startswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table_rows(lines, start_idx):
    rows = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith("|"):
        row = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        rows.append(row)
        i += 1
    return rows, i


def is_separator_row(row):
    return all(re.match(r"^:?-+:?$", cell.strip()) for cell in row if cell.strip())


def convert(md_path, docx_path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = Path(md_path).read_text(encoding="utf-8").splitlines()
    i = 0
    in_list = None  # 'ol' or 'ul'

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            in_list = None
            i += 1
            continue

        if stripped == "---":
            in_list = None
            i += 1
            continue

        if stripped.startswith("# "):
            in_list = None
            p = doc.add_heading(stripped[2:], level=0)
            i += 1
            continue

        if stripped.startswith("## "):
            in_list = None
            doc.add_heading(stripped[3:], level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            in_list = None
            doc.add_heading(stripped[4:], level=2)
            i += 1
            continue

        if stripped.startswith("|"):
            in_list = None
            rows, i = parse_table_rows(lines, i)
            if len(rows) >= 2 and is_separator_row(rows[1]):
                header, body = rows[0], rows[2:]
            else:
                header, body = rows[0], rows[1:]
            cols = len(header)
            table = doc.add_table(rows=1 + len(body), cols=cols)
            table.style = "Table Grid"
            for c, cell in enumerate(header):
                table.rows[0].cells[c].text = cell
                for run in table.rows[0].cells[c].paragraphs[0].runs:
                    run.bold = True
            for r, row in enumerate(body):
                for c in range(cols):
                    table.rows[r + 1].cells[c].text = row[c] if c < len(row) else ""
            doc.add_paragraph()
            continue

        ol_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if ol_match:
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, ol_match.group(2))
            in_list = "ol"
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, stripped[2:])
            in_list = "ul"
            i += 1
            continue

        in_list = None
        p = doc.add_paragraph()
        add_rich_text(p, stripped)
        i += 1

    doc.save(docx_path)


if __name__ == "__main__":
    src = sys.argv[1] if len(sys.argv) > 1 else "docs/mobile-deterministic-algorithms.md"
    dst = sys.argv[2] if len(sys.argv) > 2 else str(Path.home() / "Downloads" / "mobile-deterministic-algorithms.docx")
    convert(src, dst)
    print(dst)
